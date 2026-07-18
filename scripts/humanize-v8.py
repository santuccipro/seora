#!/usr/bin/env python3
"""
humanize-v8.py — Workflow ciblé (fast detect → humanise seulement les zones IA → slow rescore → PDF)
Usage: python3 humanize-v8.py input.docx output.docx [--lang fr] [--threshold 25]
       python3 humanize-v8.py input.pdf  output.docx [--lang fr]
"""

import sys, os, json, re, time, argparse, subprocess, tempfile
import urllib.request, urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document

# ─── Config ──────────────────────────────────────────────────────────────────

def load_env():
    env = {}
    for path in ["/tmp/seora-env.tmp", os.path.expanduser("~/.seora-env")]:
        try:
            with open(path) as f:
                for line in f:
                    line = line.strip()
                    if "=" in line and not line.startswith("#"):
                        k, v = line.split("=", 1)
                        env[k] = v.strip('"\'')
        except FileNotFoundError:
            pass
    return env

ENV = load_env()
DETECTOR_URL   = ENV.get("SEORA_DETECTOR_URL",  os.environ.get("SEORA_DETECTOR_URL",  "https://detector.tryseora.com"))
DETECTOR_TOKEN = ENV.get("SEORA_DETECTOR_TOKEN", os.environ.get("SEORA_DETECTOR_TOKEN", ""))

TARGET_SCORE       = 15    # % cible
CHUNK_SIZE         = 33000 # chars par appel détecteur
MIN_PARA_LEN       = 40    # chars min pour humaniser un paragraphe
REWRITE_BATCH_SIZE = 15    # paragraphes par appel claude
PARALLEL_WORKERS   = 4     # workers parallèles pour humanisation
CLAUDE_TIMEOUT     = 120
MAX_RETRIES        = 2

# ─── PDF → DOCX ──────────────────────────────────────────────────────────────

def pdf_to_docx(pdf_path: str, docx_path: str):
    from pdf2docx import Converter
    print("  Conversion PDF → DOCX...")
    c = Converter(pdf_path)
    c.convert(docx_path)
    c.close()

# ─── DOCX helpers ────────────────────────────────────────────────────────────

def get_para_text(para) -> str:
    return "".join(r.text for r in para.runs)

def set_para_text(para, new_text: str):
    if not para.runs:
        return
    new_text = new_text.replace("\n", " ").strip()
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""

def extract_paragraphs(path: str):
    doc = Document(path)
    paras = []
    for para in doc.paragraphs:
        txt = get_para_text(para).strip()
        if txt:
            paras.append((para, txt))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = get_para_text(para).strip()
                    if txt:
                        paras.append((para, txt))
    return doc, paras

# ─── Detector ────────────────────────────────────────────────────────────────

def detect_chunk(text: str, lang: str, fast_mode: bool) -> dict:
    payload = json.dumps({"text": text, "language": lang, "fast_mode": fast_mode}).encode()
    headers = {"Content-Type": "application/json"}
    if DETECTOR_TOKEN:
        headers["Authorization"] = f"Bearer {DETECTOR_TOKEN}"
    req = urllib.request.Request(
        f"{DETECTOR_URL}/detect",
        data=payload,
        headers=headers,
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=300) as r:
        return json.loads(r.read())

def detect_full(text: str, lang: str, fast_mode: bool, label: str = "") -> dict:
    chunks = []
    pos = 0
    while pos < len(text):
        end = min(pos + CHUNK_SIZE, len(text))
        if end < len(text):
            split_at = text.rfind("\n", pos + int(CHUNK_SIZE * 0.8), end)
            end = split_at if split_at > pos else end
        chunks.append(text[pos:end])
        pos = end + 1
    chunks = [c for c in chunks if c.strip()]

    pfx = f"[{label}] " if label else ""
    print(f"  {pfx}{len(chunks)} chunk(s), fast_mode={fast_mode}")

    chunk_results = []
    for i, chunk in enumerate(chunks):
        for attempt in range(3):
            try:
                r = detect_chunk(chunk, lang, fast_mode)
                chunk_results.append((chunk, r))
                print(f"  {pfx}Chunk {i+1}/{len(chunks)} → {r.get('score_global', 0):.1f}%")
                break
            except Exception as e:
                if attempt < 2:
                    time.sleep(8)
                else:
                    print(f"  {pfx}Chunk {i+1} FAIL: {e}")
                    chunk_results.append((chunk, {"score_global": 50, "zones": [], "signals": {}}))

    total_chars = sum(len(c) for c, _ in chunk_results) or 1
    merged_score = sum(r.get("score_global", 0) * len(c) / total_chars for c, r in chunk_results)
    all_zones = []
    for _, r in chunk_results:
        all_zones.extend(r.get("zones", []))
    merged_signals = {}
    for _, r in chunk_results:
        for k, v in (r.get("signals") or {}).items():
            if isinstance(v, (int, float)):
                merged_signals[k] = max(merged_signals.get(k, 0), v)

    return {
        "score_global": round(merged_score, 1),
        "signals": merged_signals,
        "zones": all_zones,
        "chunk_results": chunk_results,
    }

# ─── Ciblage des paragraphes IA ──────────────────────────────────────────────

def find_ai_paragraphs(paras: list, chunk_results: list, threshold: float) -> set:
    """Identifie les indices de paragraphes dans les chunks au-dessus du seuil."""
    para_offsets = []
    pos = 0
    for _, txt in paras:
        para_offsets.append(pos)
        pos += len(txt) + 1

    chunk_pos = 0
    high = set()
    for chunk_text, result in chunk_results:
        score = result.get("score_global", 0)
        chunk_end = chunk_pos + len(chunk_text)
        if score > threshold:
            for i, (p_start, (_, txt)) in enumerate(zip(para_offsets, paras)):
                if len(txt) < MIN_PARA_LEN:
                    continue
                p_end = p_start + len(txt)
                if p_start < chunk_end and p_end > chunk_pos:
                    high.add(i)
        chunk_pos = chunk_end + 1

    return high

# ─── Prompt builder ──────────────────────────────────────────────────────────

def build_signals_context(signals: dict) -> str:
    rules = []
    def sig(k, d=0):
        v = signals.get(k, d)
        return v if isinstance(v, (int, float)) else d

    if sig("em_dash_overuse_score") > 20:
        rules.append("• PAS de tirets longs (—) — utilise virgules ou parenthèses.")
    if sig("typographic_quotes_score") > 20:
        rules.append("• PAS de guillemets typographiques (« ») — guillemets droits ou rien.")
    if sig("connector_overuse") > 20:
        rules.append("• SUPPRIME : 'de plus', 'par ailleurs', 'ainsi', 'en outre', 'néanmoins', 'à travers', 'constitue', 'traduit'. Reformule directement.")
    if sig("tripartite_score") > 20:
        rules.append("• ÉVITE les listes de 3 éléments — utilise 2 ou 4, ou reformule en prose.")
    if sig("near_zero_typos_score") > 30:
        rules.append("• AJOUTE des marqueurs oraux : 'Franchement,', 'En fait,', 'Bref,', 'Concrètement,', 'Du coup,'")
    if sig("human_markers", 10) < 8:
        rules.append("• MARQUEURS HUMAINS insuffisants — commence 2-3 phrases par 'Je', 'Mon', 'Pour moi'.")
    if sig("burstiness", 1.0) < 0.7:
        rules.append("• VARIATION insuffisante — intercale des phrases très courtes (5-10 mots) entre les longues.")
    if sig("sentence_length_mean") > 28:
        rules.append(f"• PHRASES TROP LONGUES ({sig('sentence_length_mean'):.0f} mots en moyenne) — coupe au-dessus de 30 mots.")
    ai_hits = signals.get("ai_favorite_top") or []
    if ai_hits:
        banned = ", ".join(f'"{h.split(" ×")[0]}"' for h in ai_hits[:5])
        rules.append(f"• MOTS SUREXPLOITÉS à varier : {banned}.")
    if sig("semantic_classifier_score") > 30:
        rules.append("• STYLE trop formel/IA — commence 2-3 phrases par 'Je' ou un détail concret vécu.")
    if not rules:
        rules.append("• Ajoute des marqueurs oraux, varie la structure, rends le texte direct et naturel.")
    return "\n".join(rules)

# ─── Rewriter ────────────────────────────────────────────────────────────────

def rewrite_batch(paragraphs: list, signals_ctx: str, lang: str, offset: int = 0) -> list:
    numbered = "\n\n".join(f"[{offset+i+1}]\n{p}" for i, p in enumerate(paragraphs))
    prompt = f"""Réécris ces {len(paragraphs)} passages en {lang} pour qu'ils semblent écrits par un vrai humain. \
Ce sont des extraits de DPP (dossier de pratiques professionnelles) d'un étudiant en alternance.

RÈGLES (toutes obligatoires) :
{signals_ctx}

CONTRAINTES :
• Conserve EXACTEMENT le sens, faits, chiffres, noms, dates
• Garde le registre professionnel mais casse le style IA
• Chaque texte réécrit = une seule ligne (pas de \\n)
• Réponds UNIQUEMENT avec un JSON array, rien d'autre

FORMAT : [{{"i": 1, "t": "texte réécrit"}}, ...]

Passages :

{numbered}"""

    for attempt in range(MAX_RETRIES + 1):
        try:
            r = subprocess.run(
                ["claude", "--print", "--model", "claude-sonnet-4-6"],
                input=prompt, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT
            )
            out = r.stdout.strip()
            m = re.search(r'\[\s*\{.*?\}\s*\]', out, re.DOTALL)
            if not m:
                if attempt < MAX_RETRIES:
                    time.sleep(5)
                    continue
                return list(paragraphs)
            items = json.loads(m.group(0))
            result = list(paragraphs)
            for item in items:
                idx = item.get("i", 0) - offset - 1
                t = item.get("t", "")
                if 0 <= idx < len(result) and t and len(t) > 5:
                    result[idx] = t.replace("\n\n", " ").replace("\n", " ").strip()
            return result
        except subprocess.TimeoutExpired:
            if attempt < MAX_RETRIES:
                continue
            return list(paragraphs)
        except Exception as e:
            if attempt < MAX_RETRIES:
                time.sleep(5)
                continue
            print(f"  [FAIL] batch {offset}: {e}")
            return list(paragraphs)
    return list(paragraphs)

def rewrite_targeted(paras: list, target_indices: list, signals_ctx: str, lang: str) -> list:
    if not target_indices:
        return list(paras)
    texts = [paras[i][1] for i in target_indices]
    batches = [(i, texts[i:i+REWRITE_BATCH_SIZE]) for i in range(0, len(texts), REWRITE_BATCH_SIZE)]
    n_workers = min(PARALLEL_WORKERS, len(batches))
    print(f"  {len(batches)} batch(es) × ≤{REWRITE_BATCH_SIZE} paras ({n_workers} workers)")

    rewrites = list(texts)
    with ThreadPoolExecutor(max_workers=n_workers) as ex:
        futures = {
            ex.submit(rewrite_batch, batch, signals_ctx, lang, off): (off, len(batch))
            for off, batch in batches
        }
        for fut in as_completed(futures):
            off, blen = futures[fut]
            result = fut.result()
            changed = sum(1 for j in range(blen) if result[j] != texts[off + j])
            for j, rw in enumerate(result):
                rewrites[off + j] = rw
            print(f"  Batch [{off}..{off+blen-1}] ✓ {changed}/{blen} modifiés")

    current = list(paras)
    for j, i in enumerate(target_indices):
        para_obj, orig = paras[i]
        new_t = rewrites[j] if j < len(rewrites) else orig
        if new_t and new_t != orig and len(new_t) > 5:
            set_para_text(para_obj, new_t)
            current[i] = (para_obj, new_t)
    return current

# ─── PDF export ──────────────────────────────────────────────────────────────

def docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Converti DOCX → PDF via textutil (HTML) + Chrome headless."""
    import shutil

    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tmp:
        html_path = tmp.name

    try:
        # DOCX → HTML via textutil (natif macOS, fidèle)
        r = subprocess.run(
            ["textutil", "-convert", "html", docx_path, "-output", html_path],
            capture_output=True, text=True, timeout=30
        )
        if r.returncode != 0:
            print(f"  textutil error: {r.stderr[:100]}")
            return False

        # HTML → PDF via Chrome headless (propre, cohérent avec extraction texte)
        chrome_paths = [
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
            "/Applications/Chromium.app/Contents/MacOS/Chromium",
        ]
        chrome = next((p for p in chrome_paths if os.path.exists(p)), None)
        if not chrome:
            print("  Chrome absent, skip PDF")
            return False

        r2 = subprocess.run(
            [chrome, "--headless=new", "--disable-gpu",
             f"--print-to-pdf={pdf_path}", "--print-to-pdf-no-header",
             f"file://{html_path}"],
            capture_output=True, text=True, timeout=60
        )
        return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1000
    except Exception as e:
        print(f"  PDF error: {e}")
        return False
    finally:
        try:
            os.unlink(html_path)
        except Exception:
            pass

# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="humanize-v8 — workflow ciblé")
    parser.add_argument("input",  help="Fichier source (DOCX ou PDF)")
    parser.add_argument("output", help="Fichier de sortie (DOCX)")
    parser.add_argument("--lang",      default="fr")
    parser.add_argument("--threshold", type=float, default=25.0,
                        help="Seuil fast_mode pour cibler les zones IA (défaut: 25%%)")
    parser.add_argument("--no-pdf", action="store_true", help="Ne pas générer le PDF")
    args = parser.parse_args()

    t0 = time.time()
    print(f"\n{'='*60}")
    print(f"  humanize-v8 | {args.input} → {args.output}")
    print(f"  Seuil: {args.threshold}% | Langue: {args.lang}")
    print(f"{'='*60}")

    # 0. PDF → DOCX si besoin
    docx_input = args.input
    if args.input.lower().endswith(".pdf"):
        docx_input = args.output.replace(".docx", "_src.docx")
        pdf_to_docx(args.input, docx_input)

    # 1. Extraction
    print(f"\n[1/4] Parsing DOCX...")
    doc, paras = extract_paragraphs(docx_input)
    all_text = "\n".join(t for _, t in paras)
    print(f"  {len(paras)} paragraphes | {len(all_text)} chars")

    # 2. Fast detect → ciblage
    print(f"\n[2/4] Fast detect (rapide)...")
    t1 = time.time()
    fast_result = detect_full(all_text, args.lang, fast_mode=True, label="fast")
    fast_score = fast_result["score_global"]
    print(f"  Score fast : {fast_score:.1f}% ({time.time()-t1:.0f}s)")

    if fast_score <= TARGET_SCORE:
        print(f"  ✅ Déjà sous {TARGET_SCORE}% en fast ! On vérifie en slow...")
    else:
        ai_indices = find_ai_paragraphs(paras, fast_result["chunk_results"], args.threshold)
        signals_ctx = build_signals_context(fast_result.get("signals") or {})
        print(f"  Zones IA ciblées : {len(ai_indices)}/{len(paras)} paragraphes ({len(ai_indices)/max(len(paras),1)*100:.0f}%)")
        print(f"  Signaux:\n{signals_ctx}")

        # 3. Humanisation ciblée
        print(f"\n[3/4] Humanisation ciblée...")
        t2 = time.time()
        paras = rewrite_targeted(paras, sorted(ai_indices), signals_ctx, args.lang)
        print(f"  Humanisation : {time.time()-t2:.0f}s")

    # 4. Slow rescore
    print(f"\n[4/4] Slow rescore (score réel)...")
    t3 = time.time()
    new_text = "\n".join(t for _, t in paras)
    slow_result = detect_full(new_text, args.lang, fast_mode=False, label="slow")
    final_score = slow_result["score_global"]
    print(f"  Score final : {final_score:.1f}% ({time.time()-t3:.0f}s)")

    # Passe 2 si encore trop haut
    if final_score > TARGET_SCORE:
        print(f"\n  ⚠️ Score encore à {final_score:.1f}% — passe 2 ciblée...")
        ai2 = find_ai_paragraphs(paras, slow_result["chunk_results"], TARGET_SCORE + 5)
        signals_ctx2 = build_signals_context(slow_result.get("signals") or {})
        print(f"  Passe 2 : {len(ai2)} paragraphes ciblés")
        if ai2:
            paras = rewrite_targeted(paras, sorted(ai2), signals_ctx2, args.lang)
            new_text2 = "\n".join(t for _, t in paras)
            r2 = detect_full(new_text2, args.lang, fast_mode=False, label="slow-p2")
            final_score = r2["score_global"]
            print(f"  Score passe 2 : {final_score:.1f}%")

    # Sauvegarde DOCX
    doc.save(args.output)
    print(f"\n  💾 DOCX sauvegardé : {args.output}")

    # Export PDF
    if not args.no_pdf:
        pdf_path = args.output.replace(".docx", ".pdf")
        print(f"  Génération PDF...")
        ok = docx_to_pdf(args.output, pdf_path)
        if ok:
            print(f"  📄 PDF : {pdf_path}")
        else:
            print(f"  ⚠️  PDF non généré (erreur conversion)")

    dt = time.time() - t0
    status = "✅ OBJECTIF ATTEINT" if final_score <= TARGET_SCORE else "⚠️ encore au-dessus"
    print(f"\n{'='*60}")
    print(f"  Score FINAL : {final_score:.1f}%  {status}")
    print(f"  Durée       : {dt:.0f}s ({dt/60:.1f} min)")
    print(f"{'='*60}\n")

if __name__ == "__main__":
    main()
