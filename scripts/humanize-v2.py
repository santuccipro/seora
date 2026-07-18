#!/usr/bin/env python3
"""
humanize-v2.py — Humaniseur premium v2
Pipeline : Préparation → Détection IA → Polissage visuel → Réécriture → Validation → Export
Usage:
  python3 humanize-v2.py input.pdf   output.docx [--lang fr] [--chat-id 5002951272]
  python3 humanize-v2.py input.docx  output.docx [--lang fr] [--chat-id 5002951272]
  python3 humanize-v2.py input.pdf  --detect-only  (rapport de score uniquement)
"""

import sys, os, json, re, time, argparse, subprocess, tempfile
import urllib.request, urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Config ──────────────────────────────────────────────────────────────────

BOT_TOKEN = "8477719669:AAGgtf3fHHvoruXDs1pEu-bftcozddJMhnw"

def load_env():
    env = {}
    for path in ["/tmp/seora-env.tmp", os.path.expanduser("~/.seora-env")]:
        try:
            with open(path) as f:
                for line in f:
                    line = line.strip()
                    if "=" in line and not line.startswith("#"):
                        k, v = line.split("=", 1)
                        env[k] = v.strip('"\'')
        except FileNotFoundError:
            continue
    return env

ENV = load_env()
DETECTOR_URL   = ENV.get("SEORA_DETECTOR_URL",  os.environ.get("SEORA_DETECTOR_URL",  ""))
DETECTOR_TOKEN = ENV.get("SEORA_DETECTOR_TOKEN", os.environ.get("SEORA_DETECTOR_TOKEN", ""))

TARGET_SCORE       = 15
MAX_PASSES         = 3
CHUNK_SIZE         = 33000
MIN_PARA_LEN       = 40
REWRITE_BATCH_SIZE = 8
PARALLEL_WORKERS   = 3
CLAUDE_TIMEOUT     = 120
MAX_RETRIES        = 2

CHAT_ID = None  # set by argparse

# ─── Telegram notifications ───────────────────────────────────────────────────

def tg(msg: str):
    if not CHAT_ID or not BOT_TOKEN:
        return
    try:
        data = json.dumps({"chat_id": CHAT_ID, "text": msg}).encode()
        req = urllib.request.Request(
            f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage",
            data=data,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        urllib.request.urlopen(req, timeout=10)
    except Exception:
        pass

def log(msg: str, tg_msg: str = None):
    print(msg)
    sys.stdout.flush()
    if tg_msg:
        tg(tg_msg)

# ─── PDF → DOCX ──────────────────────────────────────────────────────────────

def pdf_to_docx(pdf_path: str, docx_path: str):
    from pdf2docx import Converter
    c = Converter(pdf_path)
    c.convert(docx_path)
    c.close()

# ─── Detector ────────────────────────────────────────────────────────────────

def detect_chunk(text: str, lang: str = "fr", fast_mode: bool = True) -> dict:
    payload = json.dumps({"text": text, "language": lang, "fast_mode": fast_mode}).encode()
    req = urllib.request.Request(
        f"{DETECTOR_URL}/detect",
        data=payload,
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {DETECTOR_TOKEN}"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=300) as r:
        return json.loads(r.read())

def split_chunks(text: str) -> list:
    chunks, pos, n = [], 0, len(text)
    while pos < n:
        end = min(pos + CHUNK_SIZE, n)
        if end < n:
            split_at = text.rfind("\n", pos + int(CHUNK_SIZE * 0.8), end)
            if split_at == -1:
                split_at = end
        else:
            split_at = n
        chunks.append(text[pos:split_at])
        pos = split_at + 1
    return [c for c in chunks if c.strip()]

def detect_all(text: str, lang: str = "fr", fast_mode: bool = True) -> dict:
    chunks = split_chunks(text)
    n = len(chunks)
    results = []
    for i, chunk in enumerate(chunks):
        for attempt in range(3):
            try:
                r = detect_chunk(chunk, lang, fast_mode)
                results.append((chunk, r))
                score = r.get("score_global", 0)
                print(f"  Chunk {i+1}/{n}: {score:.1f}%")
                sys.stdout.flush()
                break
            except Exception as e:
                if attempt < 2:
                    time.sleep(10)
                else:
                    print(f"  Chunk {i+1}/{n}: ÉCHEC ({e})")
                    results.append((chunk, {"score_global": 50, "zones": [], "signals": {}}))
    total_chars = sum(len(c) for c, _ in results)
    score = sum(r.get("score_global", 0) * len(c) / total_chars for c, r in results) if total_chars else 0
    signals = {}
    for _, r in results:
        for k, v in (r.get("signals") or {}).items():
            if isinstance(v, (int, float)):
                signals[k] = max(signals.get(k, 0), v)
    return {
        "score_global": round(score, 1),
        "signals": signals,
        "zones": [z for _, r in results for z in r.get("zones", [])],
        "chunk_results": results,
    }

def sig(signals: dict, key: str):
    v = signals.get(key, 0)
    return v if isinstance(v, (int, float)) else 0

# ─── Direct fixes + Connector swaps (zéro token) ────────────────────────────

# Connecteurs IA → alternatives naturelles (ordre : plus long d'abord pour éviter sous-match)
CONNECTOR_SWAPS = [
    (r'\bIl convient de noter que\b', 'À noter que'),
    (r'\bil convient de noter que\b', 'à noter que'),
    (r'\bIl convient de souligner que\b', 'À souligner :'),
    (r'\bil convient de souligner que\b', 'à souligner :'),
    (r'\bIl convient de\b', 'Il faut'),
    (r'\bil convient de\b', 'il faut'),
    (r'\bDans cette optique\b', 'Dans ce but'),
    (r'\bdans cette optique\b', 'dans ce but'),
    (r'\bEn ce sens\b', 'Ainsi'),
    (r'\ben ce sens\b', 'ainsi'),
    (r'\bDans le cadre de\b', 'Dans'),
    (r'\bdans le cadre de\b', 'dans'),
    (r'\bAu sein de\b', 'Dans'),
    (r'\bau sein de\b', 'dans'),
    (r'\bPar ailleurs\b', "D'ailleurs"),
    (r'\bpar ailleurs\b', "d'ailleurs"),
    (r'\bEn outre\b', 'De plus'),
    (r'\ben outre\b', 'de plus'),
    (r'\bNéanmoins\b', 'Pourtant'),
    (r'\bnéanmoins\b', 'pourtant'),
    (r'\bToutefois\b', 'Cependant'),
    (r'\btoutefois\b', 'cependant'),
    (r'\bDe plus\b', 'Aussi'),
    (r'\bde plus\b', 'aussi'),
    (r'\bEn effet\b', 'Car'),
    (r'\ben effet\b', 'car'),
    (r'\bDès lors\b', 'Ainsi'),
    (r'\bdès lors\b', 'ainsi'),
    (r'\bIl est important de\b', 'Il faut'),
    (r'\bil est important de\b', 'il faut'),
    (r'\bIl est essentiel de\b', 'Il faut'),
    (r'\bil est essentiel de\b', 'il faut'),
    (r'\bPermettre de\b', 'Permettre de'),
    (r'\bConstitue\b', 'Représente'),  # swapper l'un par l'autre pour varier
    (r"\bL'ensemble de\b", 'Les'),
    (r"\bl'ensemble de\b", 'les'),
    (r"\bL'ensemble des\b", 'Les'),
    (r"\bl'ensemble des\b", 'les'),
]

def fix_text(text: str) -> str:
    text = text.replace('—', ',').replace('–', '-')
    text = text.replace('« ', '"').replace(' »', '"')
    text = text.replace('«', '"').replace('»', '"')
    text = re.sub(r',\s*,', ',', text)
    text = re.sub(r'  +', ' ', text)
    return text.strip()

def apply_connector_swaps(text: str) -> str:
    for pattern, replacement in CONNECTOR_SWAPS:
        text = re.sub(pattern, replacement, text)
    return text

# ─── DOCX helpers ────────────────────────────────────────────────────────────

def get_text(para) -> str:
    return "".join(run.text for run in para.runs)

def set_text(para, new_text: str):
    if not para.runs:
        return
    new_text = new_text.replace("\n", " ").strip()
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""

def get_all_paras(doc) -> list:
    """Returns list of (para, text) for body paragraphs and table cells."""
    paras = []
    for para in doc.paragraphs:
        txt = get_text(para).strip()
        if txt:
            paras.append((para, txt))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = get_text(para).strip()
                    if txt:
                        paras.append((para, txt))
    return paras

# ─── ÉTAPE 3 : Polissage visuel ───────────────────────────────────────────────

def _set_cell_border(cell, **kwargs):
    """Add borders to a table cell via OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def fix_tables(doc) -> int:
    """Add borders to tables that have no visible borders. Returns count fixed."""
    fixed = 0
    for table in doc.tables:
        has_borders = False
        try:
            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None and len(tblBorders) > 0:
                    has_borders = True
        except Exception:
            pass
        if not has_borders:
            for row in table.rows:
                for cell in row.cells:
                    _set_cell_border(cell)
            fixed += 1
    return fixed

def fix_excess_empty_paras(doc) -> int:
    """Remove runs of 3+ consecutive empty paragraphs, leaving max 1."""
    consecutive = 0
    removed = 0
    paras_to_remove = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == "":
            consecutive += 1
            if consecutive > 1:
                paras_to_remove.append(para)
                removed += 1
        else:
            consecutive = 0
    for para in paras_to_remove:
        p = para._element
        p.getparent().remove(p)
    return removed

def fix_heading_spacing(doc) -> int:
    """Ensure headings have space_before so sections breathe."""
    fixed = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or style_name.startswith("Titre"):
            pf = para.paragraph_format
            if pf.space_before is None or pf.space_before.pt < 8:
                pf.space_before = Pt(10)
                fixed += 1
    return fixed

def polish_docx(doc) -> dict:
    """Run all automatic visual polishing. Returns stats."""
    stats = {}
    stats["tables_bordées"] = fix_tables(doc)
    stats["paras_vides_supprimés"] = fix_excess_empty_paras(doc)
    stats["titres_espacés"] = fix_heading_spacing(doc)
    return stats

# ─── ÉTAPE 4 : Réécriture anti-IA ────────────────────────────────────────────

def build_rules(signals: dict) -> str:
    rules = []
    if sig(signals, "em_dash_overuse_score") > 20:
        rules.append("• N'utilise PAS de tirets longs (—). Remplace par virgule ou parenthèses.")
    if sig(signals, "typographic_quotes_score") > 20:
        rules.append('• N\'utilise PAS « guillemets fr ». Utilise des guillemets droits " ".')
    if sig(signals, "connector_overuse") > 20:
        rules.append("• SUPPRIME connecteurs IA : 'de plus', 'par ailleurs', 'ainsi', 'en outre', 'néanmoins', 'il convient de'. Reformule directement.")
    if sig(signals, "tripartite_score") > 20:
        rules.append("• ÉVITE les triades (premièrement/deuxièmement/troisièmement). Utilise 2 ou 4 éléments, ou formule en prose.")
    if sig(signals, "near_zero_typos_score") > 30:
        rules.append("• Ajoute un marqueur oral occasionnel ('Franchement,', 'Concrètement,', 'Du coup,'). Varie les longueurs de phrases.")
    ai_hits = signals.get("ai_favorite_top") or []
    if ai_hits:
        banned = ", ".join(f'"{h.split(" ×")[0]}"' for h in ai_hits[:5])
        rules.append(f"• ÉVITE ces mots/phrases surexploités par l'IA : {banned}.")
    if sig(signals, "semantic_classifier_score") > 30:
        rules.append("• Commence certaines phrases par 'Je' ou un détail concret. Brise les schémas formels.")
    if not rules:
        rules.append("• Ajoute des marqueurs naturels, varie la structure, rends le texte direct.")
    return "\n".join(rules)

def rewrite_batch(paragraphs: list, rules: str, lang: str, offset: int = 0) -> list:
    items = [{"i": offset + i + 1, "t": p} for i, p in enumerate(paragraphs)]
    prompt = f"""Tu es un expert en rédaction professionnelle. Réécris ces passages en {lang} pour qu'ils sonnent HUMAINS.

RÈGLES OBLIGATOIRES :
{rules}

CONTRAINTES :
• Conserve EXACTEMENT le sens, les faits, les chiffres, les noms propres
• Garde le registre professionnel adapté à un DPP étudiant
• Chaque texte réécrit = une seule ligne (pas de \\n dans le texte)

RÉPONSE : JSON array UNIQUEMENT, sans aucune explication :
[{{"i": 1, "t": "texte réécrit"}}, ...]

Passages à réécrire :
{json.dumps(items, ensure_ascii=False)}"""

    for attempt in range(MAX_RETRIES + 1):
        try:
            r = subprocess.run(
                ["claude", "--print", "--model", "claude-sonnet-4-6"],
                input=prompt, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT
            )
            out = r.stdout.strip()
            match = re.search(r'\[\s*\{.*?\}\s*\]', out, re.DOTALL)
            if not match:
                if attempt < MAX_RETRIES:
                    time.sleep(5); continue
                return list(paragraphs)
            rewrites = json.loads(match.group(0))
            result = list(paragraphs)
            for item in rewrites:
                idx = item.get("i", 0) - offset - 1
                if 0 <= idx < len(result):
                    t = item.get("t", "").replace("\n\n", " ").replace("\n", " ").strip()
                    if t and len(t) > 5:
                        result[idx] = t
            return result
        except subprocess.TimeoutExpired:
            if attempt < MAX_RETRIES:
                print(f"  [retry] offset={offset} timeout"); continue
            return list(paragraphs)
        except Exception as e:
            if attempt < MAX_RETRIES:
                time.sleep(5); continue
            print(f"  [FAIL] offset={offset}: {e}")
            return list(paragraphs)
    return list(paragraphs)

def rewrite_all(paras_to_rewrite: list, rules: str, lang: str) -> list:
    batches = [(i, paras_to_rewrite[i:i+REWRITE_BATCH_SIZE])
               for i in range(0, len(paras_to_rewrite), REWRITE_BATCH_SIZE)]
    n_workers = min(PARALLEL_WORKERS, len(batches))
    print(f"  {len(batches)} batches × ≤{REWRITE_BATCH_SIZE} paras ({n_workers} workers)")
    sys.stdout.flush()
    result = list(paras_to_rewrite)
    with ThreadPoolExecutor(max_workers=n_workers) as ex:
        futures = {
            ex.submit(rewrite_batch, batch, rules, lang, offset): (offset, len(batch))
            for offset, batch in batches
        }
        for fut in as_completed(futures):
            offset, blen = futures[fut]
            rw = fut.result()
            changed = sum(1 for j in range(blen) if rw[j] != paras_to_rewrite[offset+j])
            for j, t in enumerate(rw):
                result[offset+j] = t
            print(f"  [{offset}..{offset+blen-1}] ✓ {changed}/{blen} réécrits")
            sys.stdout.flush()
    return result

def find_high_paras(all_paras: list, chunk_results: list, threshold: float) -> list:
    offsets, pos = [], 0
    for _, txt in all_paras:
        offsets.append(pos)
        pos += len(txt) + 1
    chunk_pos = 0
    high = set()
    for chunk_text, result in chunk_results:
        score = result.get("score_global", 0)
        chunk_end = chunk_pos + len(chunk_text)
        if score > threshold:
            for i, (p_off, (_, txt)) in enumerate(zip(offsets, all_paras)):
                if len(txt) < MIN_PARA_LEN: continue
                if p_off < chunk_end and p_off + len(txt) > chunk_pos:
                    high.add(i)
        chunk_pos = chunk_end + 1
    return sorted(high)

# ─── Export PDF (LibreOffice headless) ───────────────────────────────────────

def export_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert DOCX → PDF via LibreOffice headless. Returns True if success."""
    out_dir = os.path.dirname(pdf_path) or "/tmp"
    try:
        r = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", out_dir],
            capture_output=True, text=True, timeout=120
        )
        # soffice saves as <basename>.pdf in out_dir
        base = os.path.splitext(os.path.basename(docx_path))[0]
        generated = os.path.join(out_dir, base + ".pdf")
        if os.path.exists(generated):
            os.rename(generated, pdf_path)
            return True
        return False
    except Exception as e:
        print(f"  [PDF] LibreOffice failed: {e}")
        return False

# ─── MODE DÉTECTION SEULE ────────────────────────────────────────────────────

def mode_detect_only(input_path: str, lang: str):
    log("📊 MODE DÉTECTION SEULE", f"📊 Détection en cours sur {os.path.basename(input_path)}...")

    if input_path.lower().endswith(".pdf"):
        import fitz
        doc = fitz.open(input_path)
        text = "".join(page.get_text() for page in doc)
        doc.close()
    else:
        doc = Document(input_path)
        text = "\n".join(get_text(p).strip() for p in doc.paragraphs if get_text(p).strip())

    print(f"  {len(text)} chars, détection Seora...")
    result = detect_all(text, lang, fast_mode=True)
    score = result["score_global"]
    signals = result.get("signals", {})
    top_signals = [(k, v) for k, v in signals.items() if isinstance(v, (int, float)) and v > 30]
    top_signals.sort(key=lambda x: -x[1])

    report = f"""📋 RAPPORT DE DÉTECTION IA
Fichier : {os.path.basename(input_path)}
Score global : {score:.1f}% {"✅ OK" if score < 15 else "⚠️ IA détectée"}

Signaux principaux (>30%) :
""" + "\n".join(f"  • {k}: {v:.0f}%" for k, v in top_signals[:8])

    log(report, report)

# ─── MODE HUMANISATION ───────────────────────────────────────────────────────

def mode_humanize(input_path: str, output_docx: str, lang: str, export_pdf_too: bool = False):
    t0 = time.time()
    fname = os.path.basename(input_path)

    # ─── ÉTAPE 1 — PRÉPARATION ───────────────────────────────────────────────
    log(f"\n{'='*60}\nÉTAPE 1/5 — PRÉPARATION\n{'='*60}",
        f"⚙️ Étape 1/5 — Préparation\n{fname}")

    docx_src = input_path
    if input_path.lower().endswith(".pdf"):
        docx_src = output_docx.replace(".docx", "_src.docx")
        log(f"  PDF → DOCX ({os.path.basename(input_path)})...")
        pdf_to_docx(input_path, docx_src)
        log(f"  Converti en {os.path.basename(docx_src)}")

    doc = Document(docx_src)
    all_paras = get_all_paras(doc)
    log(f"  {len(all_paras)} paragraphes extraits")

    # Fixes directs + swaps connecteurs
    n_fixed = 0
    n_swaps = 0
    for para_obj, txt in all_paras:
        fixed = fix_text(txt)
        swapped = apply_connector_swaps(fixed)
        if swapped != txt:
            set_text(para_obj, swapped)
            if fixed != txt:
                n_fixed += 1
            if swapped != fixed:
                n_swaps += 1
    # Refresh texts after fixes
    all_paras = get_all_paras(doc)
    log(f"  {n_fixed} fixes directs (tirets, guillemets) + {n_swaps} swaps connecteurs IA")

    # ─── ÉTAPE 2 — DÉTECTION IA ──────────────────────────────────────────────
    log(f"\n{'='*60}\nÉTAPE 2/5 — DÉTECTION IA\n{'='*60}",
        f"🔍 Étape 2/5 — Détection IA en cours...")

    full_text = "\n".join(txt for _, txt in all_paras)
    chunks = split_chunks(full_text)
    log(f"  {len(full_text)} chars → {len(chunks)} chunks")

    detect_result = detect_all(full_text, lang, fast_mode=True)
    init_score = detect_result["score_global"]
    signals = detect_result.get("signals", {})
    chunk_results = detect_result.get("chunk_results", [])

    top_sigs = sorted(
        [(k, v) for k, v in signals.items() if isinstance(v, (int, float)) and v > 25],
        key=lambda x: -x[1]
    )[:6]
    sigs_str = ", ".join(f"{k.replace('_score','')}: {v:.0f}%" for k, v in top_sigs)
    log(f"  Score initial : {init_score:.1f}%\n  Signaux : {sigs_str}",
        f"🔍 Score initial : {init_score:.1f}%\nSignaux : {sigs_str}")

    if init_score <= TARGET_SCORE:
        log(f"  ✅ Déjà sous {TARGET_SCORE}% — pas de réécriture nécessaire")
        tg(f"✅ Document déjà à {init_score:.1f}% — export direct.")
        doc.save(output_docx)
        log(f"\n✅ Output : {output_docx}")
        return

    # ─── ÉTAPE 3 — POLISSAGE VISUEL ──────────────────────────────────────────
    log(f"\n{'='*60}\nÉTAPE 3/5 — POLISSAGE VISUEL\n{'='*60}",
        f"✨ Étape 3/5 — Polissage visuel...")

    polish_stats = polish_docx(doc)
    stats_str = ", ".join(f"{k}: {v}" for k, v in polish_stats.items() if v > 0)
    if stats_str:
        log(f"  Corrections : {stats_str}", f"✨ Polissage : {stats_str}")
    else:
        log(f"  Document déjà propre visuellement")
    # Refresh paragraphs after polishing
    all_paras = get_all_paras(doc)

    # ─── ÉTAPE 4 — RÉÉCRITURE ANTI-IA ────────────────────────────────────────
    log(f"\n{'='*60}\nÉTAPE 4/5 — RÉÉCRITURE ANTI-IA\n{'='*60}",
        f"✍️ Étape 4/5 — Réécriture Claude Sonnet 4.6...")

    rules = build_rules(signals)
    current_paras = list(all_paras)
    current_score = init_score
    prev_chunk_results = chunk_results

    for pass_num in range(1, MAX_PASSES + 1):
        if current_score <= TARGET_SCORE:
            break

        MAX_CLAUDE_PARAS = 20  # limit Claude rewrites to avoid re-generating IA text en masse

        if pass_num == 1:
            # Target longest paragraphs first (most likely to contain IA patterns)
            candidates = [(i, len(txt)) for i, (_, txt) in enumerate(current_paras) if len(txt) >= MIN_PARA_LEN]
            candidates.sort(key=lambda x: -x[1])
            target_idx = [i for i, _ in candidates[:MAX_CLAUDE_PARAS]]
            target_idx.sort()
            log(f"\n  Pass {pass_num} — {len(target_idx)} paras ciblés (top {MAX_CLAUDE_PARAS} les plus longs)...")
        else:
            target_idx = find_high_paras(current_paras, prev_chunk_results, TARGET_SCORE + 5)
            if not target_idx:
                log(f"  Pass {pass_num} : aucun para IA ciblé — stop")
                break
            log(f"\n  Pass {pass_num} — {len(target_idx)} paras IA ciblés...")
            tg(f"✍️ Pass {pass_num} — {len(target_idx)} paras restants à réécrire...")

        t_rw = time.time()
        target_texts = [current_paras[i][1] for i in target_idx]
        rewrites = rewrite_all(target_texts, rules, lang)
        n_changed = 0
        for j, i in enumerate(target_idx):
            para_obj, orig = current_paras[i]
            new_t = rewrites[j] if j < len(rewrites) else orig
            if new_t and new_t != orig and len(new_t) > 5:
                set_text(para_obj, new_t)
                current_paras[i] = (para_obj, new_t)
                n_changed += 1
        log(f"  {n_changed}/{len(target_idx)} paras réécrits en {time.time()-t_rw:.0f}s")

        # ─── ÉTAPE 5 — VALIDATION ─────────────────────────────────────────────
        log(f"\n{'='*60}\nÉTAPE 5/5 — VALIDATION SCORE\n{'='*60}",
            f"📐 Étape 5/5 — Validation score (pass {pass_num})...")

        new_text = "\n".join(txt for _, txt in current_paras)
        rescore = detect_all(new_text, lang, fast_mode=(pass_num < MAX_PASSES))
        prev_chunk_results = rescore.get("chunk_results", [])
        current_score = rescore.get("score_global", 0)
        signals = rescore.get("signals", {})
        rules = build_rules(signals)

        status = "✅" if current_score <= TARGET_SCORE else "⚠️"
        log(f"  Score après pass {pass_num} : {current_score:.1f}% {status}",
            f"📐 Score après pass {pass_num} : {current_score:.1f}% {status}")

        if current_score <= TARGET_SCORE:
            break

    # ─── ÉTAPE 6 — EXPORT ────────────────────────────────────────────────────
    log(f"\n{'='*60}\nÉTAPE 6 — EXPORT\n{'='*60}",
        f"💾 Étape 6 — Export en cours...")

    doc.save(output_docx)
    dt = int(time.time() - t0)
    status_txt = "✅ OBJECTIF ATTEINT" if current_score <= TARGET_SCORE else "⚠️ encore > 15%"

    log(f"\n{'='*60}")
    log(f"  Score AVANT  : {init_score:.1f}%")
    log(f"  Score APRÈS  : {current_score:.1f}%  {status_txt}")
    log(f"  Durée totale : {dt}s ({dt//60}min {dt%60}s)")
    log(f"  Output DOCX  : {output_docx}")

    tg(f"""✅ Humanisation terminée !

Score avant : {init_score:.1f}%
Score après : {current_score:.1f}% {status_txt}
Durée : {dt//60}min {dt%60}s
Fichier : {os.path.basename(output_docx)}""")

    if export_pdf_too:
        log(f"\n  Conversion PDF (LibreOffice)...")
        pdf_path = output_docx.replace(".docx", ".pdf")
        ok = export_pdf(output_docx, pdf_path)
        if ok:
            log(f"  PDF : {pdf_path}")
            tg(f"📄 PDF prêt : {os.path.basename(pdf_path)}")
        else:
            log(f"  PDF : échec LibreOffice")
            tg("⚠️ PDF conversion échouée — DOCX disponible")

# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    global CHAT_ID
    parser = argparse.ArgumentParser()
    parser.add_argument("input", help="Input PDF or DOCX")
    parser.add_argument("output", nargs="?", default=None, help="Output DOCX path")
    parser.add_argument("--lang", default="fr")
    parser.add_argument("--detect-only", action="store_true")
    parser.add_argument("--chat-id", default="5002951272", help="Telegram chat_id for notifications")
    parser.add_argument("--pdf", action="store_true", help="Also export PDF via LibreOffice")
    args = parser.parse_args()

    CHAT_ID = args.chat_id

    if args.detect_only:
        mode_detect_only(args.input, args.lang)
    else:
        if not args.output:
            base = os.path.splitext(args.input)[0]
            args.output = base + "_humanized.docx"
        mode_humanize(args.input, args.output, args.lang, export_pdf_too=args.pdf)

if __name__ == "__main__":
    main()
