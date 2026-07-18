#!/usr/bin/env python3
"""
humanize-fast.py v7 — Humaniseur premium (Claude CLI local, rapide)
Usage: python3 humanize-fast.py input.pdf output.docx [--lang fr]
       python3 humanize-fast.py input.docx output.docx [--lang fr]

Flow:
 0. Si PDF : conversion → DOCX (pdf2docx)
 1. Parsing DOCX → paragraphes
 2. Remplacements directs (em-dash, guillemets, mots IA) — instant
 3. Détection fast_mode → score + signaux
 4. Réécriture avec `claude --print` LOCAL (rapide, pas de tunnel)
 5. Rescoring slow_mode → vérification ≤15%
 6. Pass 2 ciblé si besoin
"""

import sys, os, json, re, time, argparse, subprocess, tempfile
import urllib.request, urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document

# ─── Config ──────────────────────────────────────────────────────────────────

def load_env():
    env = {}
    try:
        with open("/tmp/seora-env.tmp") as f:
            for line in f:
                line = line.strip()
                if "=" in line and not line.startswith("#"):
                    k, v = line.split("=", 1)
                    env[k] = v.strip('"\'')
    except FileNotFoundError:
        pass
    return env

ENV = load_env()
DETECTOR_URL   = ENV.get("SEORA_DETECTOR_URL",  os.environ.get("SEORA_DETECTOR_URL",  ""))
DETECTOR_TOKEN = ENV.get("SEORA_DETECTOR_TOKEN", os.environ.get("SEORA_DETECTOR_TOKEN", ""))

TARGET_GLOBAL      = 15    # target IA score
MAX_PASSES         = 3
CHUNK_SIZE         = 33000 # chars per detector call
MIN_PARA_LEN       = 30    # min chars to rewrite
REWRITE_BATCH_SIZE = 10    # paragraphs per claude --print call
PARALLEL_WORKERS   = 3     # parallel claude --print calls
CLAUDE_TIMEOUT     = 120   # seconds per claude call
MAX_RETRIES        = 2

# ─── PDF → DOCX ──────────────────────────────────────────────────────────────

def pdf_to_docx(pdf_path: str, docx_path: str):
    from pdf2docx import Converter
    print(f"  Conversion PDF → DOCX...")
    c = Converter(pdf_path)
    c.convert(docx_path, start=0, end=None)
    c.close()
    print(f"  Converti : {docx_path}")

# ─── Direct string fixes ─────────────────────────────────────────────────────

# Mots IA → variantes humaines (liste cyclée selon position dans le texte)
AI_WORD_VARIANTS = [
    # (pattern, [variantes]) — la variante est choisie par hash du contexte
    (r'\bconstitue\b',          ['est', 'forme', 'représente', 'est']),
    (r'\bconstituent\b',        ['sont', 'forment', 'représentent', 'sont']),
    (r'\bcontribue\b',          ['aide', 'participe à', 'joue un rôle dans']),
    (r'\bcontribuent\b',        ['aident', 'participent à', 'jouent un rôle dans']),
    (r'\btraduit\b',            ['montre', 'reflète', 'illustre']),
    (r'\btraduisent\b',         ['montrent', 'reflètent', 'illustrent']),
    (r'\brévèle\b',             ['montre', 'indique', 'dévoile']),
    (r'\brévèlent\b',           ['montrent', 'indiquent', 'dévoilent']),
    (r'\bs\'inscrit dans\b',    ['fait partie de', 's\'intègre dans', 'relève de']),
    (r'\bs\'inscrivent dans\b', ['font partie de', 's\'intègrent dans', 'relèvent de']),
    (r'\bdans le cadre de\b',   ['dans', 'pour', 'lors de']),
    (r'\bde plus\b',            ['aussi', 'et', 'en plus']),
    (r'\bpar ailleurs\b',       ['aussi', 'en outre', 'de même']),
    (r'\ben outre\b',           ['et aussi', 'de même', 'en plus']),
    (r'\bainsi\b',              ['donc', 'du coup', 'alors']),
    (r'\bnéanmoins\b',          ['mais', 'cependant', 'toutefois']),
    (r'\bpermettant de\b',      ['aidant à', 'servant à', 'donnant la possibilité de']),
    (r'\bpermet de\b',          ['aide à', 'sert à', 'donne la possibilité de']),
    (r'\bpermettent de\b',      ['aident à', 'servent à', 'donnent la possibilité de']),
    (r'\bmettre en place\b',    ['créer', 'instaurer', 'lancer']),
    (r'\bà travers\b',          ['via', 'par', 'grâce à']),
    (r'\bau regard de\b',       ['selon', 'par rapport à', 'vu']),
    (r'\bs\'avère\b',           ['est', 'se révèle', 'apparaît']),
    (r'\bs\'avèrent\b',         ['sont', 'se révèlent', 'apparaissent']),
]

_subst_counter = {}

def apply_word_substitutions(text: str) -> str:
    """Remplace les mots IA typiques par des variantes humaines, cyclées."""
    for pattern, variants in AI_WORD_VARIANTS:
        def replace_fn(m, p=pattern, vs=variants):
            _subst_counter[p] = _subst_counter.get(p, 0) + 1
            return vs[(_subst_counter[p] - 1) % len(vs)]
        text = re.sub(pattern, replace_fn, text, flags=re.IGNORECASE)
    return text

def apply_direct_fixes(text: str) -> str:
    text = text.replace('—', ',')
    text = text.replace('–', '-')
    text = text.replace('« ', '"').replace(' »', '"')
    text = text.replace('« ', '"').replace(' »', '"')
    text = text.replace('«', '"').replace('»', '"')
    text = re.sub(r',\s*,', ',', text)
    text = re.sub(r'  +', ' ', text)
    text = apply_word_substitutions(text)
    return text.strip()

def apply_direct_fixes_to_paras(paras: list) -> tuple:
    changed = 0
    new_paras = []
    for para_obj, txt in paras:
        fixed = apply_direct_fixes(txt)
        if fixed != txt:
            set_para_text(para_obj, fixed)
            new_paras.append((para_obj, fixed))
            changed += 1
        else:
            new_paras.append((para_obj, txt))
    return new_paras, changed

# ─── Detector ────────────────────────────────────────────────────────────────

def detect_chunk(text: str, lang: str = "fr", fast_mode: bool = True) -> dict:
    payload = json.dumps({"text": text, "language": lang, "fast_mode": fast_mode}).encode()
    req = urllib.request.Request(
        f"{DETECTOR_URL}/detect",
        data=payload,
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {DETECTOR_TOKEN}"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=300) as r:
        return json.loads(r.read())

def split_text_chunks(text: str) -> list:
    chunks = []
    pos = 0
    n = len(text)
    while pos < n:
        end = min(pos + CHUNK_SIZE, n)
        if end == n:
            chunks.append(text[pos:end])
            break
        split_at = text.rfind("\n", pos + int(CHUNK_SIZE * 0.8), end)
        if split_at == -1:
            split_at = end
        chunks.append(text[pos:split_at])
        pos = split_at + 1
    return [c for c in chunks if c.strip()]

def detect(text: str, lang: str = "fr", fast_mode: bool = True, label: str = "") -> dict:
    chunks = split_text_chunks(text)
    n = len(chunks)
    pfx = f"[{label}] " if label else ""
    print(f"  {pfx}{n} chunk(s) (fast_mode={fast_mode})")

    chunk_results = []
    for i, chunk in enumerate(chunks):
        for attempt in range(3):
            try:
                r = detect_chunk(chunk, lang, fast_mode)
                chunk_results.append((chunk, r))
                score = r.get("score_global", 0)
                print(f"  {pfx}Chunk {i+1}/{n} ✓ ({len(chunk)} chars, {score:.1f}%)")
                break
            except Exception as e:
                if attempt < 2:
                    time.sleep(10)
                else:
                    print(f"  {pfx}Chunk {i+1} ÉCHEC")
                    chunk_results.append((chunk, {"score_global": 50, "zones": [], "signals": {}}))

    total_chars = sum(len(c) for c, _ in chunk_results)
    merged_score = sum(r.get("score_global", 0) * len(c) / total_chars for c, r in chunk_results) if total_chars else 0
    all_zones = []
    for _, r in chunk_results:
        all_zones.extend(r.get("zones", []))
    merged_signals = {}
    for _, r in chunk_results:
        for k, v in (r.get("signals") or {}).items():
            if isinstance(v, (int, float)):
                merged_signals[k] = max(merged_signals.get(k, 0), v)

    return {
        "score_global": round(merged_score, 1),
        "signals": merged_signals,
        "zones": all_zones,
        "chunk_results": chunk_results,
    }

# ─── Position-based HIGH para matching ───────────────────────────────────────

def find_high_para_indices_by_position(paras: list, chunk_results: list, threshold: float = 20.0) -> list:
    para_offsets = []
    pos = 0
    for _, txt in paras:
        para_offsets.append(pos)
        pos += len(txt) + 1

    chunk_pos = 0
    high_indices = set()
    for chunk_text, result in chunk_results:
        score = result.get("score_global", 0)
        chunk_end = chunk_pos + len(chunk_text)
        if score > threshold:
            for i, (p_start, (_, txt)) in enumerate(zip(para_offsets, paras)):
                if len(txt) < MIN_PARA_LEN:
                    continue
                p_end = p_start + len(txt)
                if p_start < chunk_end and p_end > chunk_pos:
                    high_indices.add(i)
        chunk_pos = chunk_end + 1

    return sorted(high_indices)

# ─── Signal helpers ───────────────────────────────────────────────────────────

def sig(signals: dict, key: str, default=0):
    v = signals.get(key, default)
    return v if isinstance(v, (int, float)) else default

def build_signals_context(result: dict) -> str:
    signals = result.get("signals") or {}
    rules = []
    if sig(signals, "em_dash_overuse_score") > 20:
        rules.append("• NE METS PAS de tirets longs (—). Utilise des virgules ou parenthèses.")
    if sig(signals, "typographic_quotes_score") > 20:
        rules.append('• NE METS PAS guillemets typographiques (« »). Utilise des guillemets droits ou rien.')
    if sig(signals, "connector_overuse") > 20:
        rules.append("• SUPPRIME les connecteurs IA : 'de plus', 'par ailleurs', 'ainsi', 'en outre', 'néanmoins', 'à travers', 'constitue', 'traduit'. Reformule directement sans eux.")
    if sig(signals, "tripartite_score") > 20:
        rules.append("• ÉVITE les listes de 3 éléments. Utilise 2 ou 4 éléments, ou reformule en prose.")
    if sig(signals, "near_zero_typos_score") > 30:
        rules.append("• AJOUTE au moins 3 marqueurs oraux/directs dans le texte : 'Franchement,', 'En fait,', 'Bref,', 'Concrètement,', 'Du coup,', 'Je dois dire que'. Place-les en début de phrase.")
    human_markers = sig(signals, "human_markers", 10)
    if human_markers < 8:
        rules.append(f"• MARQUEURS HUMAINS INSUFFISANTS ({int(human_markers)} détectés). Commence 2-3 phrases par 'Je', 'Mon', 'Pour moi' ou une observation personnelle courte (5-8 mots) avant de développer.")
    burstiness = sig(signals, "burstiness", 1.0)
    if burstiness < 0.7:
        rules.append(f"• VARIATION DE LONGUEUR TROP FAIBLE (burstiness {burstiness:.2f}). Intercale des phrases très courtes (5-10 mots) entre les phrases longues. Exemple : 'C'est le cas ici.' ou 'Voilà le problème.'")
    sentence_len = sig(signals, "sentence_length_mean", 0)
    if sentence_len > 28:
        rules.append(f"• PHRASES TROP LONGUES (moyenne {sentence_len:.0f} mots). Coupe les phrases de plus de 30 mots en 2 phrases distinctes. Vise 15-22 mots en moyenne.")
    ai_hits = signals.get("ai_favorite_top") or []
    if ai_hits:
        banned = ", ".join(f'"{h.split(" ×")[0]}"' for h in ai_hits[:5])
        rules.append(f"• MOTS SUREXPLOITÉS à remplacer par des synonymes variés : {banned}.")
    if sig(signals, "semantic_classifier_score") > 30:
        rules.append("• STYLE TROP FORMEL/IA. Commence 2-3 phrases par 'Je' ou un détail concret vécu. Brise les schémas formels abstraits.")
    if sig(signals, "rhetorical_score") > 15:
        rules.append("• ÉVITE les structures rhétoriques en 3 temps (primo/secundo/tertio, D'une part/D'autre part/Enfin).")
    if not rules:
        rules.append("• Ajoute des marqueurs oraux, varie la structure, rends le texte direct et naturel.")
    return "\n".join(rules)

# ─── DOCX helpers ────────────────────────────────────────────────────────────

def get_para_text(para) -> str:
    return "".join(run.text for run in para.runs)

def set_para_text(para, new_text: str):
    if not para.runs:
        return
    new_text = new_text.replace("\n", " ").strip()
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""

def extract_docx_paragraphs(path: str):
    doc = Document(path)
    paras = []
    for para in doc.paragraphs:
        txt = get_para_text(para).strip()
        if txt:
            paras.append((para, txt))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = get_para_text(para).strip()
                    if txt:
                        paras.append((para, txt))
    return doc, paras

# ─── Claude CLI rewriter (local, fast) ───────────────────────────────────────

def claude_rewrite_one_batch(paragraphs: list, signals_context: str, lang: str, offset: int = 0) -> list:
    """Rewrite via local `claude --print`. No network tunnel needed."""
    numbered = "\n\n".join(f"[{offset+i+1}]\n{p}" for i, p in enumerate(paragraphs))

    prompt = f"""Tu es un rédacteur expert qui humanise des textes pour tromper les détecteurs d'IA. Réécris ces {len(paragraphs)} passages en {lang} pour qu'ils semblent écrits par un vrai humain.

RÈGLES CRITIQUES (toutes obligatoires) :
{signals_context}

CONTRAINTES ABSOLUES :
• Conserve EXACTEMENT le sens, les faits, les chiffres, les noms propres, les dates
• Garde le registre professionnel mais CASSE le style IA
• Varie activement la longueur des phrases (courtes ET longues)
• Chaque texte réécrit = une seule ligne (pas de \\n dans le texte)
• N'ajoute pas d'explication, de commentaire ou de préambule

RÉPONSE : JSON array UNIQUEMENT, sans aucun texte avant ou après :
[{{"i": 1, "t": "texte réécrit"}}, {{"i": 2, "t": "texte réécrit"}}, ...]

Passages à réécrire :

{numbered}"""

    for attempt in range(MAX_RETRIES + 1):
        try:
            result = subprocess.run(
                ["claude", "--print", "--model", "claude-sonnet-4-6"],
                input=prompt,
                capture_output=True,
                text=True,
                timeout=CLAUDE_TIMEOUT
            )
            text_out = result.stdout.strip()
            match = re.search(r'\[\s*\{.*?\}\s*\]', text_out, re.DOTALL)
            if not match:
                if attempt < MAX_RETRIES:
                    time.sleep(5)
                    continue
                return list(paragraphs)
            rewrites_raw = json.loads(match.group(0))
            output = list(paragraphs)
            for item in rewrites_raw:
                idx = item.get("i", 0) - offset - 1
                if 0 <= idx < len(output):
                    t = item.get("t", "")
                    if t and len(t) > 5:
                        t = t.replace("\n\n", " ").replace("\n", " ").strip()
                        output[idx] = t
            return output
        except subprocess.TimeoutExpired:
            if attempt < MAX_RETRIES:
                print(f"  [retry] batch {offset}: timeout, retry {attempt+1}...")
                continue
            print(f"  [FAIL] batch {offset}: timeout après {MAX_RETRIES+1} tentatives")
            return list(paragraphs)
        except Exception as e:
            if attempt < MAX_RETRIES:
                time.sleep(5)
                continue
            print(f"  [FAIL] batch {offset}: {e}")
            return list(paragraphs)

    return list(paragraphs)

def claude_rewrite_batch(paragraphs: list, signals_context: str, lang: str = "fr") -> list:
    if not paragraphs:
        return []
    batches = [(i, paragraphs[i:i+REWRITE_BATCH_SIZE]) for i in range(0, len(paragraphs), REWRITE_BATCH_SIZE)]
    n_workers = min(PARALLEL_WORKERS, len(batches))
    print(f"  {len(batches)} batches × ≤{REWRITE_BATCH_SIZE} paras ({n_workers} workers parallèles)")
    result = list(paragraphs)

    with ThreadPoolExecutor(max_workers=n_workers) as ex:
        futures = {
            ex.submit(claude_rewrite_one_batch, batch, signals_context, lang, offset): (offset, len(batch))
            for offset, batch in batches
        }
        for fut in as_completed(futures):
            offset, blen = futures[fut]
            rewrites = fut.result()
            changed = sum(1 for j in range(blen) if rewrites[j] != paragraphs[offset+j])
            for j, rw in enumerate(rewrites):
                result[offset + j] = rw
            print(f"  [{offset}..{offset+blen-1}] ✓ {changed}/{blen}")

    return result

# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input",  help="Input PDF or DOCX path")
    parser.add_argument("output", help="Output DOCX path")
    parser.add_argument("--lang", default="fr")
    args = parser.parse_args()

    t0 = time.time()
    print(f"\n[humanize-fast v7] {args.input} → {args.output}")

    # 0. PDF → DOCX if needed
    docx_input = args.input
    if args.input.lower().endswith(".pdf"):
        docx_input = args.output.replace(".docx", "_src.docx")
        pdf_to_docx(args.input, docx_input)

    # 1. Extract
    print(f"\n[1/5] Parsing DOCX...")
    doc, all_paras = extract_docx_paragraphs(docx_input)
    print(f"  {len(all_paras)} paragraphes")

    # 2. Direct fixes
    print(f"\n[2/5] Remplacements directs...")
    all_paras, n_direct = apply_direct_fixes_to_paras(all_paras)
    print(f"  {n_direct} modifiés")

    # 3. Slow detection (mode précis — toujours)
    all_text = "\n".join(txt for _, txt in all_paras)
    print(f"\n[3/5] Détection slow_mode ({len(all_text)} chars)...")
    t_d = time.time()
    init_result = detect(all_text, args.lang, fast_mode=False, label="slow")
    init_score = init_result.get("score_global", 0)
    print(f"  Score initial : {init_score:.1f}% ({time.time()-t_d:.0f}s)")

    if init_score <= TARGET_GLOBAL:
        print(f"  ✅ Déjà sous {TARGET_GLOBAL}% !")
        doc.save(args.output)
        return

    signals_ctx = build_signals_context(init_result)
    print(f"  Signaux :\n{signals_ctx}")

    # 4. Rewrite passes
    subst_indices = [i for i, (_, txt) in enumerate(all_paras) if len(txt) >= MIN_PARA_LEN]
    current_paras = list(all_paras)
    current_score = init_score
    prev_chunk_results = init_result.get("chunk_results", [])

    for pass_num in range(1, MAX_PASSES + 1):
        if current_score <= TARGET_GLOBAL:
            break

        if pass_num == 1:
            target_indices = subst_indices
            print(f"\n[4/5] Pass {pass_num} — {len(target_indices)} paras (≥{MIN_PARA_LEN} chars)...")
        else:
            target_indices = find_high_para_indices_by_position(
                current_paras, prev_chunk_results, threshold=TARGET_GLOBAL + 5
            )
            if not target_indices:
                print(f"\n  Pass {pass_num} : aucun para HIGH, stop")
                break
            print(f"\n[4/5] Pass {pass_num} — {len(target_indices)} paras HIGH ciblés...")

        target_texts = [current_paras[i][1] for i in target_indices]
        t_rw = time.time()
        rewrites = claude_rewrite_batch(target_texts, signals_ctx, args.lang)
        dt_rw = time.time() - t_rw

        changed = 0
        for j, i in enumerate(target_indices):
            para_obj, orig = current_paras[i]
            new_t = rewrites[j] if j < len(rewrites) else orig
            if new_t and new_t != orig and len(new_t) > 5:
                set_para_text(para_obj, new_t)
                current_paras[i] = (para_obj, new_t)
                changed += 1
        print(f"  {changed}/{len(target_indices)} changés en {dt_rw:.0f}s")

        # Rescore
        print(f"\n[5/5] Rescoring pass {pass_num} (slow mode)...")
        new_text = "\n".join(txt for _, txt in current_paras)
        t_r = time.time()
        rescore = detect(new_text, args.lang, fast_mode=False, label=f"p{pass_num}")
        prev_chunk_results = rescore.get("chunk_results", [])
        current_score = rescore.get("score_global", 0)
        print(f"  Score : {current_score:.1f}% ({time.time()-t_r:.0f}s)")
        signals_ctx = build_signals_context(rescore)

        if current_score <= TARGET_GLOBAL:
            print(f"  ✅ Objectif atteint !")
            break

    final_score = current_score
    doc.save(args.output)
    dt = time.time() - t0
    status = "✅ OBJECTIF ATTEINT" if final_score <= TARGET_GLOBAL else "⚠️ encore au-dessus"
    print(f"\n{'='*60}")
    print(f"  Score AVANT  : {init_score:.1f}%")
    print(f"  Score APRÈS  : {final_score:.1f}%  {status}")
    print(f"  Durée totale : {dt:.0f}s ({dt/60:.1f}min)")
    print(f"  Output       : {args.output}")

if __name__ == "__main__":
    main()
