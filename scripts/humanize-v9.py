#!/usr/bin/env python3
"""
humanize-v9.py — Workflow humanisation pérenne
- Préserve 100% la mise en page
- Input/output dans le même format (PDF→PDF, DOCX→DOCX)
- 1 seule passe LLM sur les zones IA uniquement (économe en tokens)
- <10 min sur un DPP standard

Usage:
  python3 humanize-v9.py input.pdf   [--lang fr] [--threshold 25]
  python3 humanize-v9.py input.docx  [--lang fr] [--threshold 25]

Output: même dossier que l'input, suffixe _humanise
"""

import sys, os, json, re, time, argparse, subprocess, tempfile
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document

# ─── Config ──────────────────────────────────────────────────────────────────

def load_env():
    env = {}
    for path in ["/tmp/seora-env.tmp", os.path.expanduser("~/.seora-env")]:
        try:
            with open(path) as f:
                for line in f:
                    line = line.strip()
                    if "=" in line and not line.startswith("#"):
                        k, v = line.split("=", 1)
                        env[k] = v.strip('"\'')
        except FileNotFoundError:
            pass
    return env

ENV = load_env()
DETECTOR_URL   = ENV.get("SEORA_DETECTOR_URL",  "https://detector.tryseora.com")
DETECTOR_TOKEN = ENV.get("SEORA_DETECTOR_TOKEN", "")

TARGET_SCORE   = 15.0   # % cible slow_mode
CHUNK_SIZE     = 30000  # chars par appel détecteur
MIN_PARA_LEN   = 50     # chars min pour humaniser
BATCH_SIZE     = 12     # paragraphes par appel claude
WORKERS        = 4      # workers parallèles
CLAUDE_TIMEOUT = 120
MAX_RETRIES    = 2

CHROME_PATHS = [
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    "/Applications/Chromium.app/Contents/MacOS/Chromium",
]

# Patterns IA à supprimer directement (sans LLM)
DIRECT_FIXES = [
    (r" — ", ", "), (r"— ", ", "), (r" —", ","), (r"–", "-"),
    (r"«\s*", '"'), (r"\s*»", '"'),
    (r"\bde plus\b", "aussi"), (r"\bpar ailleurs\b", "aussi"),
    (r"\ben outre\b", "et"), (r"\bainsi que\b", "et"),
    (r"\bnéanmoins\b", "mais"), (r"\bcependant\b", "mais"),
    (r"\btoutefois\b", "mais"), (r"\ben effet\b", ""),
    (r"\bà travers\b", "via"), (r"\bau regard de\b", "selon"),
    (r"\bdans le cadre de\b", "dans"),
    (r"\bconstitue\b", "est"), (r"\bconstituent\b", "sont"),
    (r"\bs'avère\b", "est"), (r"\bs'avèrent\b", "sont"),
    (r"\bpermet de\b", "aide à"), (r"\bpermettent de\b", "aident à"),
    (r"\bcontribue à\b", "aide à"),
]

# ─── DOCX helpers — préservation mise en page ────────────────────────────────

def get_para_text(para) -> str:
    return "".join(r.text for r in para.runs)

def set_para_text_preserve_layout(para, new_text: str):
    """
    Remplace le texte d'un paragraphe en préservant la mise en page.
    - Si 1 run : remplace directement
    - Si N runs : met le nouveau texte dans le 1er run (style conservé),
      vide les suivants SAUF s'ils sont bold/italic (dans ce cas on garde
      leur formatage et on répartit le texte proportionnellement)
    """
    if not para.runs:
        return

    new_text = new_text.replace("\n\n", " ").replace("\n", " ").strip()

    # Cas simple : 1 seul run
    if len(para.runs) == 1:
        para.runs[0].text = new_text
        return

    # Cas complexe : plusieurs runs
    # Vérifier si des runs ont du bold/italic/underline
    has_formatting = any(
        r.bold or r.italic or r.underline
        for r in para.runs[1:]
        if r.text.strip()
    )

    if not has_formatting:
        # Pas de formatage inline critique → tout dans le 1er run
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        # Il y a du formatage inline → répartition proportionnelle
        orig_total = sum(len(r.text) for r in para.runs) or 1
        pos = 0
        for i, run in enumerate(para.runs):
            if i == len(para.runs) - 1:
                run.text = new_text[pos:]
            else:
                ratio = len(run.text) / orig_total
                end = pos + max(1, int(len(new_text) * ratio))
                # Couper sur un espace pour ne pas couper un mot
                while end < len(new_text) and new_text[end] != " ":
                    end += 1
                run.text = new_text[pos:end]
                pos = end

def apply_direct_fixes(text: str) -> str:
    """Substitutions directes — déterministes, pas de LLM."""
    for pat, repl in DIRECT_FIXES:
        text = re.sub(pat, repl, text, flags=re.IGNORECASE)
    text = re.sub(r"[​‌‍­﻿⁠]", "", text)   # zero-width chars
    text = re.sub(r";([a-zA-ZÀ-ÿ])", r"\1", text)  # point-virgule parasite
    text = re.sub(r"  +", " ", text)
    text = re.sub(r" ,", ",", text)
    return text.strip()

# ─── Detector ────────────────────────────────────────────────────────────────

def _detect_chunk(text: str, lang: str, fast_mode: bool) -> dict:
    payload = json.dumps({"text": text, "language": lang, "fast_mode": fast_mode}).encode()
    headers = {"Content-Type": "application/json"}
    if DETECTOR_TOKEN:
        headers["Authorization"] = f"Bearer {DETECTOR_TOKEN}"
    req = urllib.request.Request(
        f"{DETECTOR_URL}/detect", data=payload, headers=headers, method="POST"
    )
    with urllib.request.urlopen(req, timeout=300) as r:
        return json.loads(r.read())

def detect(text: str, lang: str, fast_mode: bool, label: str = "") -> dict:
    chunks, pos = [], 0
    while pos < len(text):
        end = min(pos + CHUNK_SIZE, len(text))
        if end < len(text):
            split = text.rfind("\n", pos + int(CHUNK_SIZE * 0.8), end)
            end = split if split > pos else end
        chunks.append(text[pos:end])
        pos = end + 1
    chunks = [c for c in chunks if c.strip()]

    pfx = f"[{label}] " if label else ""
    chunk_results = []
    for i, chunk in enumerate(chunks):
        for attempt in range(3):
            try:
                r = _detect_chunk(chunk, lang, fast_mode)
                chunk_results.append((chunk, r))
                print(f"  {pfx}chunk {i+1}/{len(chunks)} → {r.get('score_global', 0):.1f}%")
                break
            except Exception as e:
                if attempt < 2:
                    time.sleep(8)
                else:
                    print(f"  {pfx}chunk {i+1} FAIL: {e}")
                    chunk_results.append((chunk, {"score_global": 50, "zones": [], "signals": {}}))

    total = sum(len(c) for c, _ in chunk_results) or 1
    score = sum(r.get("score_global", 0) * len(c) / total for c, r in chunk_results)
    signals = {}
    for _, r in chunk_results:
        for k, v in (r.get("signals") or {}).items():
            if isinstance(v, (int, float)):
                signals[k] = max(signals.get(k, 0), v)
    zones = []
    for _, r in chunk_results:
        zones.extend(r.get("zones") or [])

    return {"score_global": round(score, 1), "signals": signals,
            "zones": zones, "chunk_results": chunk_results}

def find_ai_paragraphs(paras: list, chunk_results: list, threshold: float) -> set:
    offsets, pos = [], 0
    for _, txt in paras:
        offsets.append(pos)
        pos += len(txt) + 1

    high, chunk_pos = set(), 0
    for chunk_text, result in chunk_results:
        chunk_end = chunk_pos + len(chunk_text)
        if result.get("score_global", 0) > threshold:
            for i, (pstart, (_, txt)) in enumerate(zip(offsets, paras)):
                if len(txt) >= MIN_PARA_LEN:
                    if pstart < chunk_end and (pstart + len(txt)) > chunk_pos:
                        high.add(i)
        chunk_pos = chunk_end + 1
    return high

# ─── Prompt builder ──────────────────────────────────────────────────────────

def build_signals_prompt(signals: dict) -> str:
    def sig(k, d=0):
        v = signals.get(k, d)
        return v if isinstance(v, (int, float)) else d

    rules = [
        "• PAS de tirets longs (—), PAS de guillemets typographiques (« »)",
        "• PAS de connecteurs IA : 'de plus', 'par ailleurs', 'ainsi', 'en outre', 'néanmoins', 'à travers', 'constitue', 'traduit', 'en effet'",
        "• PAS de listes de 3 éléments (A, B et C) — utilise 2 ou reformule en prose",
    ]
    if sig("near_zero_typos_score") > 30:
        rules.append("• AJOUTE des marqueurs oraux : 'En fait,', 'Bref,', 'Concrètement,', 'Du coup,'")
    if sig("human_markers", 10) < 8:
        rules.append("• Commence 2-3 phrases par 'Je', 'Mon', 'Pour moi'")
    if sig("burstiness", 1.0) < 0.7:
        rules.append("• VARIE la longueur : mélange phrases courtes (8-12 mots) et longues (20-25 mots)")
    if sig("sentence_length_mean") > 28:
        rules.append(f"• COUPE les phrases > 30 mots en 2")
    if sig("semantic_classifier_score") > 30:
        rules.append("• STYLE trop formel — commence par 'Je' ou un fait concret vécu")
    return "\n".join(rules)

# ─── LLM rewriter ────────────────────────────────────────────────────────────

def rewrite_batch(paragraphs: list, signals_prompt: str, lang: str, offset: int = 0) -> list:
    numbered = "\n\n".join(f"[{offset+i+1}]\n{p}" for i, p in enumerate(paragraphs))
    prompt = f"""Réécris ces {len(paragraphs)} passages en {lang} pour qu'ils sonnent écrits par un étudiant en alternance (DPP Master Gestion de Patrimoine).

RÈGLES (toutes obligatoires) :
{signals_prompt}

CONTRAINTES ABSOLUES :
• Conserve EXACTEMENT les faits, chiffres, noms propres, dates, références
• Une seule ligne par passage (pas de \\n)
• UNIQUEMENT le JSON array en réponse, rien d'autre

FORMAT : [{{"i": 1, "t": "texte réécrit"}}, ...]

{numbered}"""

    for attempt in range(MAX_RETRIES + 1):
        try:
            r = subprocess.run(
                ["claude", "--print", "--model", "claude-sonnet-4-6"],
                input=prompt, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT
            )
            m = re.search(r'\[\s*\{.*?\}\s*\]', r.stdout.strip(), re.DOTALL)
            if not m:
                if attempt < MAX_RETRIES:
                    time.sleep(5)
                    continue
                return list(paragraphs)
            items = json.loads(m.group(0))
            result = list(paragraphs)
            for item in items:
                idx = item.get("i", 0) - offset - 1
                t = item.get("t", "")
                if 0 <= idx < len(result) and t and len(t) > 5:
                    # Post-process : supprimer em-dashes réintroduits par Claude
                    t = t.replace(" — ", ", ").replace("— ", ", ").replace(" —", ",")
                    result[idx] = t.replace("\n\n", " ").replace("\n", " ").strip()
            return result
        except subprocess.TimeoutExpired:
            if attempt < MAX_RETRIES:
                continue
            return list(paragraphs)
        except Exception as e:
            if attempt < MAX_RETRIES:
                time.sleep(5)
                continue
            return list(paragraphs)
    return list(paragraphs)

def humanize_targeted(paras: list, target_idx: list, signals_prompt: str, lang: str) -> list:
    if not target_idx:
        return paras
    texts = [paras[i][1] for i in target_idx]
    batches = [(off, texts[off:off+BATCH_SIZE]) for off in range(0, len(texts), BATCH_SIZE)]
    n_workers = min(WORKERS, len(batches))
    print(f"  {len(batches)} batch(es), {n_workers} workers, {len(texts)} paras")

    rewrites = list(texts)
    with ThreadPoolExecutor(max_workers=n_workers) as ex:
        futures = {
            ex.submit(rewrite_batch, batch, signals_prompt, lang, off): (off, len(batch))
            for off, batch in batches
        }
        for fut in as_completed(futures):
            off, blen = futures[fut]
            res = fut.result()
            changed = sum(1 for j in range(blen) if res[j] != texts[off+j])
            for j, rw in enumerate(res):
                rewrites[off+j] = rw
            print(f"  [{off}..{off+blen-1}] ✓ {changed}/{blen} modifiés")

    result = list(paras)
    for j, i in enumerate(target_idx):
        para_obj, orig = paras[i]
        new_t = rewrites[j]
        if new_t and new_t != orig and len(new_t) > 5:
            set_para_text_preserve_layout(para_obj, new_t)
            result[i] = (para_obj, new_t)
    return result

# ─── PDF helpers ─────────────────────────────────────────────────────────────

def pdf_to_docx(pdf_path: str) -> str:
    from pdf2docx import Converter
    out = pdf_path.replace(".pdf", "_src.docx")
    print(f"  PDF → DOCX...")
    c = Converter(pdf_path)
    c.convert(out)
    c.close()
    return out

def docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as tmp:
        html = tmp.name
    try:
        r = subprocess.run(
            ["textutil", "-convert", "html", docx_path, "-output", html],
            capture_output=True, text=True, timeout=30
        )
        if r.returncode != 0:
            return False
        chrome = next((p for p in CHROME_PATHS if os.path.exists(p)), None)
        if not chrome:
            return False
        subprocess.run(
            [chrome, "--headless=new", "--disable-gpu",
             f"--print-to-pdf={pdf_path}", "--print-to-pdf-no-header",
             f"file://{html}"],
            capture_output=True, timeout=60
        )
        return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 1000
    except Exception as e:
        print(f"  PDF error: {e}")
        return False
    finally:
        try:
            os.unlink(html)
        except Exception:
            pass

# ─── Extract paragraphs ──────────────────────────────────────────────────────

def extract_paragraphs(doc):
    """Retourne [(para_obj, text)] — body + tableaux."""
    result = []
    for para in doc.paragraphs:
        txt = get_para_text(para).strip()
        if txt:
            result.append((para, txt))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = get_para_text(para).strip()
                    if txt:
                        result.append((para, txt))
    return result

def full_text(paras: list) -> str:
    return "\n".join(t for _, t in paras)

# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="humanize-v9 — workflow pérenne")
    parser.add_argument("input", help="Fichier source (PDF ou DOCX)")
    parser.add_argument("--lang", default="fr")
    parser.add_argument("--threshold", type=float, default=25.0,
                        help="Seuil fast_mode pour cibler zones IA (défaut: 25%%)")
    parser.add_argument("--output", help="Chemin de sortie (optionnel)")
    args = parser.parse_args()

    t0 = time.time()
    input_is_pdf = args.input.lower().endswith(".pdf")
    base = os.path.splitext(args.input)[0]
    ext = ".pdf" if input_is_pdf else ".docx"
    out_path = args.output or f"{base}_humanise{ext}"

    print(f"\n{'='*60}")
    print(f"  humanize-v9 | {os.path.basename(args.input)} → {os.path.basename(out_path)}")
    print(f"  Seuil: {args.threshold}% | Langue: {args.lang}")
    print(f"{'='*60}")

    # 0. PDF → DOCX si besoin
    if input_is_pdf:
        docx_work = pdf_to_docx(args.input)
    else:
        docx_work = args.input

    # 1. Parse
    print(f"\n[1/5] Parsing...")
    doc = Document(docx_work)
    paras = extract_paragraphs(doc)
    print(f"  {len(paras)} paragraphes | {len(full_text(paras))} chars")

    # 2. Fixes directs (déterministe, sans LLM)
    print(f"\n[2/5] Nettoyage direct (sans LLM)...")
    n_fixed = 0
    for para_obj, txt in paras:
        new = apply_direct_fixes(txt)
        if new != txt:
            set_para_text_preserve_layout(para_obj, new)
            n_fixed += 1
    paras = extract_paragraphs(doc)  # re-extraire après fixes
    print(f"  {n_fixed} paragraphes nettoyés")

    # 3. Fast detect → ciblage
    print(f"\n[3/5] Fast detect...")
    t1 = time.time()
    fast_result = detect(full_text(paras), args.lang, fast_mode=True, label="fast")
    fast_score = fast_result["score_global"]
    print(f"  Score fast : {fast_score:.1f}% ({time.time()-t1:.0f}s)")

    if fast_score <= TARGET_SCORE:
        print(f"  Déjà sous {TARGET_SCORE}% !")
        ai_idx = set()
    else:
        # Threshold adaptatif : si score fast proche du threshold, on l'abaisse
        effective_threshold = min(args.threshold, max(fast_score * 0.6, TARGET_SCORE))
        print(f"  Threshold effectif : {effective_threshold:.1f}%")
        ai_idx = find_ai_paragraphs(paras, fast_result["chunk_results"], effective_threshold)
        # Si 0 zones trouvées, forcer sur tous les paragraphes body
        if not ai_idx:
            ai_idx = {i for i, (_, txt) in enumerate(paras) if len(txt) >= MIN_PARA_LEN}
            print(f"  Aucune zone ciblée → humanise tous les paragraphes body")
        signals_prompt = build_signals_prompt(fast_result.get("signals") or {})
        print(f"  Zones IA : {len(ai_idx)}/{len(paras)} paragraphes")
        print(f"  Signaux:\n{signals_prompt}")

    # 4. Humanisation LLM ciblée
    if ai_idx:
        print(f"\n[4/5] Humanisation ciblée (LLM)...")
        t2 = time.time()
        paras = humanize_targeted(paras, sorted(ai_idx), signals_prompt, args.lang)
        print(f"  Durée : {time.time()-t2:.0f}s")
    else:
        print(f"\n[4/5] Humanisation : rien à faire")

    # 5. Slow rescore
    print(f"\n[5/5] Slow rescore...")
    t3 = time.time()
    slow_result = detect(full_text(paras), args.lang, fast_mode=False, label="slow")
    final_score = slow_result["score_global"]
    print(f"  Score FINAL : {final_score:.1f}% ({time.time()-t3:.0f}s)")

    # Passe 2 si encore trop haut (ciblage sur zones slow >TARGET+5)
    if final_score > TARGET_SCORE:
        from_slow = find_ai_paragraphs(paras, slow_result["chunk_results"], TARGET_SCORE + 5)
        if from_slow:
            print(f"\n  Passe 2 : {len(from_slow)} paragraphes slow-ciblés...")
            sp2 = build_signals_prompt(slow_result.get("signals") or {})
            paras = humanize_targeted(paras, sorted(from_slow), sp2, args.lang)
            slow2 = detect(full_text(paras), args.lang, fast_mode=False, label="slow-p2")
            final_score = slow2["score_global"]
            print(f"  Score passe 2 : {final_score:.1f}%")

    # Sauvegarder DOCX
    docx_out = out_path if not input_is_pdf else f"{base}_humanise.docx"
    doc.save(docx_out)
    print(f"\n  DOCX sauvegardé : {docx_out}")

    # Convertir en PDF si l'input était un PDF
    if input_is_pdf:
        print(f"  Conversion PDF...")
        ok = docx_to_pdf(docx_out, out_path)
        if ok:
            print(f"  PDF : {out_path}")
        else:
            print(f"  PDF non généré — livraison DOCX")
            out_path = docx_out

    dt = time.time() - t0
    status = "✅ OBJECTIF ATTEINT" if final_score <= TARGET_SCORE else f"⚠️  {final_score:.1f}% (plancher ~15% pour DPP académique)"
    print(f"\n{'='*60}")
    print(f"  Score FINAL : {final_score:.1f}%  {status}")
    print(f"  Durée       : {dt:.0f}s ({dt/60:.1f} min)")
    print(f"  Output      : {out_path}")
    print(f"{'='*60}\n")
    return out_path, final_score

if __name__ == "__main__":
    main()
